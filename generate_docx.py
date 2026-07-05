"""
Generate file .docx jadwal sholat.

CATATAN FONT: Diminta font "Segoe UI Black", tapi itu font proprietary Microsoft
yang TIDAK ADA di Linux (runner GitHub Actions jalan di Ubuntu) - kalau nama itu
dipaksa dipakai, LibreOffice akan diam-diam mengganti ke font fallback acak yang
lebih lebar, dan itu penyebab tabel jadi wrap/berantakan sebelumnya. Sebagai
gantinya dipakai "Archivo Black" - font gratis (Google Fonts, lisensi OFL) yang
memang didesain di berat "Black" (setebal Segoe UI Black), jadi hasilnya secara
visual setara, dan dijamin selalu tersedia karena diinstall langsung di workflow
GitHub Actions (lihat langkah "Install font Archivo Black" di generate-jadwal.yml).

Isi tabel:
- Header tabel: latar ungu (5F497A), teks putih, bold
- Baris data: latar biru muda (B6DDE8) merata di semua sel, teks bold
- Border tabel: garis tunggal hitam 0.75pt di semua sisi
- Highlight (di atas latar biru muda), teks jadi PUTIH biar kebaca jelas:
    * Baris Senin & Kamis  -> sel HARI + sel MAGRIB, latar DARK_BLUE
      (mengikuti sunnah puasa Senin-Kamis, highlight waktu berbuka)
    * Baris Jumat          -> sel HARI + sel ZUHUR, latar GREEN
      (mengingatkan waktu shalat Jumat)
    * Tanggal Hijriah 13/14/15 -> sel H, latar DARK_BLUE (Ayyamul Bidh)
- Baris Ahad: teks "Ahad" di sel HARI diberi warna MERAH (hari libur/weekend),
  tanpa latar highlight.
- Semua kolom di-"autofit": lebar kolom otomatis menyesuaikan isi terlebar di
  kolom itu, jadi kolom dengan isi pendek (M, H, DUHA, ASAR, ISYA) otomatis
  menyempit, dan tidak ada teks yang terpaksa turun ke baris ke-2.
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HEADER_FILL = "5F497A"
DATA_FILL = "B6DDE8"
FONT_NAME = "Archivo Black"
WHITE = "FFFFFF"
RED = "C00000"

HEADERS = ["M", "H", "HARI", "IMSAK", "SUBUH", "SYURUQ", "DUHA", "ZUHUR", "ASAR", "MAGRIB", "ISYA"]
COL_IDX = {name: i for i, name in enumerate(HEADERS)}

# Lebar tiap kolom (cm) - disesuaikan ke panjang isi terlebarnya sendiri:
# M/H cuma angka pendek -> sempit. HARI harus muat "Selasa"/"Sabtu" dalam
# 1 baris. Kolom jam disesuaikan ke panjang nama headernya (SYURUQ/MAGRIB
# 6 huruf perlu paling lebar, DUHA/ASAR/ISYA 4 huruf paling sempit).
COLUMN_WIDTHS_CM = {
    "M": 0.5, "H": 0.5, "HARI": 2.55,
    "IMSAK": 1.85, "SUBUH": 1.85, "SYURUQ": 2.15, "DUHA": 1.55,
    "ZUHUR": 1.85, "ASAR": 1.55, "MAGRIB": 2.05, "ISYA": 1.55,
}


def _set_cell_shading(cell, fill_hex: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "000000")
        borders.append(el)
    tblPr.append(borders)


def _set_table_cell_margins(table, margin_dxa: int = 20):
    """Margin sel dikecilkan supaya tidak ada ruang kosong berlebih yang
    bikin teks gampang wrap / tabel keliatan boros tempat."""
    tblPr = table._tbl.tblPr
    mar = OxmlElement("w:tblCellMar")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:w"), str(margin_dxa))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tblPr.append(mar)


def _disable_autofit(table):
    """Paksa lebar kolom fixed sesuai COLUMN_WIDTHS_CM (bukan autofit Word),
    supaya konsisten persis sama di semua penampil/converter."""
    table.autofit = False
    tblPr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tblPr.append(layout)

    total_cm = sum(COLUMN_WIDTHS_CM[h] for h in HEADERS)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:type"), "dxa")
    tblW.set(qn("w:w"), str(int(total_cm * 566.9)))

    for row in table.rows:
        for idx, header in enumerate(HEADERS):
            row.cells[idx].width = Cm(COLUMN_WIDTHS_CM[header])


def _write_cell(cell, text, *, color=None, highlight=None, size=11):
    """Semua isi sel selalu BOLD dan pakai font Archivo Black (sesuai
    permintaan tebal/black). Kalau ada highlight (DARK_BLUE/GREEN), warna
    teks otomatis dipaksa putih supaya kontras & kebaca jelas."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(str(text))
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.font.bold = True

    if highlight:
        run.font.highlight_color = highlight
        run.font.color.rgb = RGBColor.from_string(WHITE)
    elif color:
        run.font.color.rgb = RGBColor.from_string(color)


def build_document(data: dict, hijriah_range_label: str) -> Document:
    """
    data: hasil dari fetch_jadwal.get_jadwal_bulanan(), sudah ditambah key
          "h_tanggal" (int hari Hijriah) di tiap item jadwal (lihat main.py).
    hijriah_range_label: contoh "Dzulhijjah 1447 - Muharram 1448"
    """
    doc = Document()

    section = doc.sections[0]
    section.orientation = 0  # portrait
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(0.5)

    # Halaman dibuat 1 halaman besar (bukan A4 biasa) supaya tabel 28-31 baris
    # tidak pernah kepotong ke halaman 2. Kelebihan area putih di bawah akan
    # dipotong otomatis nanti di convert_to_image.py (auto-crop), jadi tidak
    # perlu kalkulasi tinggi per baris yang rapuh (rentan meleset kalau font
    # fallback bikin teks wrap 2 baris di sel HARI).
    # Kertas A4 standar. Sudah dites: dengan font Archivo Black + lebar kolom
    # yang dikalibrasi (lihat COLUMN_WIDTHS_CM), tabel 28-31 baris selalu muat
    # rapi dalam 1 halaman A4 tanpa perlu ukuran kertas custom.
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)

    # --- Judul ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Jadwal Sholat {data['kabkota']}")
    run.font.name = FONT_NAME
    run.font.size = Pt(18)
    run.font.bold = True

    # --- Sub judul: bulan/tahun + rentang hijriah + sumber ---
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = sub.add_run(f"{data['bulan_nama'].upper()} {data['tahun']} (")
    r1.font.name = FONT_NAME
    r1.font.size = Pt(13)
    r1.font.bold = True
    r2 = sub.add_run(hijriah_range_label)
    r2.font.name = FONT_NAME
    r2.font.size = Pt(10)
    r2.font.bold = True
    r2.font.underline = True
    r3 = sub.add_run(")   ")
    r3.font.name = FONT_NAME
    r3.font.size = Pt(13)
    r3.font.bold = True
    r4 = sub.add_run("Sumber: bimasislam.kemenag.go.id")
    r4.font.name = FONT_NAME
    r4.font.size = Pt(9)
    r4.font.bold = True
    r4.italic = True

    doc.add_paragraph()

    # --- Tabel ---
    n_rows = len(data["jadwal"]) + 1
    table = doc.add_table(rows=n_rows, cols=len(HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_borders(table)
    _set_table_cell_margins(table)

    # Header
    for c, label in enumerate(HEADERS):
        cell = table.rows[0].cells[c]
        _set_cell_shading(cell, HEADER_FILL)
        _write_cell(cell, label, color="FFFFFF", size=10)

    # Data
    for r, item in enumerate(data["jadwal"], start=1):
        hari = item["hari"]
        if hari in ("Minggu",):
            hari = "Ahad"  # ikut gaya penamaan Kemenag

        h_tanggal = item["h_tanggal"]

        row_values = {
            "M": item["tanggal"],
            "H": h_tanggal,
            "HARI": hari,
            "IMSAK": item["imsak"],
            "SUBUH": item["subuh"],
            "SYURUQ": item["terbit"],
            "DUHA": item["dhuha"],
            "ZUHUR": item["dzuhur"],
            "ASAR": item["ashar"],
            "MAGRIB": item["maghrib"],
            "ISYA": item["isya"],
        }

        is_senin_kamis = hari in ("Senin", "Kamis")
        is_jumat = hari == "Jumat"
        is_ayyamul_bidh = h_tanggal in (13, 14, 15)

        for col_name, col_idx in COL_IDX.items():
            cell = table.rows[r].cells[col_idx]
            _set_cell_shading(cell, DATA_FILL)

            highlight = None
            color = None
            if is_senin_kamis and col_name in ("HARI", "MAGRIB"):
                highlight = WD_COLOR_INDEX.DARK_BLUE
            if is_jumat and col_name in ("HARI", "ZUHUR"):
                highlight = WD_COLOR_INDEX.GREEN
            if is_ayyamul_bidh and col_name == "H":
                highlight = WD_COLOR_INDEX.DARK_BLUE
            if hari == "Ahad" and col_name == "HARI" and not highlight:
                color = RED

            _write_cell(cell, row_values[col_name], color=color, highlight=highlight, size=10.5)

    _disable_autofit(table)
    return doc
